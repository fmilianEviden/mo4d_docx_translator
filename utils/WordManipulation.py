import os

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches


h1_counter = 0

def check_content(content: str) -> bool:
    h1_list = ['documento técnico - ', 'document technique - ']

    content_lower = content.lower()
    for item in h1_list:
        if item in content_lower:
            return True
    return False

def add_element_to_word(content, style, doc: Document, styles_dict: dict):
    global h1_counter

    if style == 'img':
        doc = add_image_to_doc(content, doc)
    elif style == 'tab':
        doc = add_table_to_doc(content, doc)
    else:
        if content == "":
            return doc
        elif content.startswith(" Technical Document - "):
            content = content.replace(" Technical Document - ", "")
            style_name = styles_dict.get("h1")
            h1_counter += 1
            content = f"{h1_counter}. {content}"
        elif check_content(content):
            content = content.split('-')[1].strip()
            style_name = styles_dict.get("h1")
            h1_counter += 1
            content = f"{h1_counter}. {content}"
        else:
            style_name = styles_dict.get(style)

        p = doc.add_paragraph()
        run = p.add_run(content)
        style_to_apply = doc.styles[style_name]
        if style_to_apply.type == WD_STYLE_TYPE.PARAGRAPH:
            p.style = style_to_apply
        elif style_to_apply.type == WD_STYLE_TYPE.CHARACTER:
            run.style = style_to_apply

        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    return doc


def get_image(img_tag, html_file):
    if 'src' in img_tag.attrs:
        img_src = img_tag['src']
        img_filename = os.path.basename(img_src)
        img_path = os.path.join(os.path.dirname(html_file), img_filename)
        return img_path, "img"


def add_image_to_doc(img_path, doc: Document):
    """Adds an image to a Word document"""
    if not img_path or img_path.endswith('logo_icon.png'):
        return doc

    try:
        paragraph = doc.add_paragraph()
        paragraph.add_run().add_picture(img_path, width=Inches(4))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception as e:
        print(f"Error adding image to document: {e}")

    return doc


def get_table(table_tag, html_file):
    """Processes table tags to return table data as tuples"""
    rows = table_tag.find_all('tr')
    if not rows:
        return

    table_data = []
    for row in rows:
        cells = row.find_all(['th', 'td'])
        row_data = tuple(cell.get_text(strip=True) for cell in cells)
        table_data.append(row_data)

    return table_data, "tab"


def add_table_to_doc(table_data, doc: Document):
    """Adds table data to a Word document"""
    if not table_data:
        return doc

    num_rows = len(table_data)
    num_cols = len(table_data[0])
    table = doc.add_table(rows=num_rows, cols=num_cols)

    try:
        table.style = 'EvidenTableStyle'
    except Exception as e:
        print(f"Error applying table style: {e}")
        pass

    for i, row_data in enumerate(table_data):
        for j, cell_text in enumerate(row_data):
            table.cell(i, j).text = cell_text

    return doc


def get_styles_from_user_template(template_file_path: str):
    doc = Document(template_file_path)
    if not doc:
        return []

    styles = [style.name for style in doc.styles]

    return styles
